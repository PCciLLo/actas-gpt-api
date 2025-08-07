from docx import Document

def generar_acta(template_path, salida_path, datos):
    doc = Document(template_path)

    def reemplazar_campos_texto():
        for p in doc.paragraphs:
            for run in p.runs:
                for key, value in datos.items():
                    if isinstance(value, str):
                        marcador = f"{{{{{key}}}}}"
                        if marcador in run.text:
                            run.text = run.text.replace(marcador, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for key, value in datos.items():
                                if isinstance(value, str):
                                    marcador = f"{{{{{key}}}}}"
                                    if marcador in run.text:
                                        run.text = run.text.replace(marcador, value)

    def rellenar_tablas():
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{PARTICIPANTES}}" in cell.text:
                        table._tbl.remove(row._tr)
                        for item in datos["PARTICIPANTES"]:
                            nueva = table.add_row().cells
                            nueva[0].text = item["id"]
                            nueva[1].text = item["asistente"]
                            nueva[2].text = item["iniciales"]
                            nueva[3].text = item["compania"]
                    elif "{{PROXIMOS_PASOS}}" in cell.text:
                        table._tbl.remove(row._tr)
                        for paso in datos["PROXIMOS_PASOS"]:
                            nueva = table.add_row().cells
                            nueva[0].text = paso["responsable"]
                            nueva[1].text = paso["titulo"]
                            nueva[2].text = paso["descripcion"]
                            nueva[3].text = paso["fecha"]
                            nueva[4].text = paso["estado"]

    def insertar_listas():
        claves = ["TEMAS", "ACUERDOS"]
        for clave in claves:
            marcador = f"{{{{{clave}}}}}"
            items = datos.get(clave, [])
            p_idx = 0
            while p_idx < len(doc.paragraphs):
                p = doc.paragraphs[p_idx]
                if marcador in p.text:
                    parent = p._element.getparent()
                    idx = parent.index(p._element)
                    parent.remove(p._element)
                    for j, item in enumerate(items):
                        nuevo = doc.add_paragraph(item)
                        parent.insert(idx + j, nuevo._element)
                    break
                p_idx += 1

    reemplazar_campos_texto()
    rellenar_tablas()
    insertar_listas()
    doc.save(salida_path)
